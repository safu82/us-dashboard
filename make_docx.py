#!/usr/bin/env python3
"""Render a newsletter markdown draft into a nicely formatted Word (.docx).

Handles the markdown our synthesis produces: # title, ### dek, ## sections,
--- rules, GFM tables, - bullets, and inline **bold** / *italic*. Styles it
with a branded title, navy headings, shaded table headers, and readable body
type — something you can send as-is or lightly edit.

Usage:
  python make_docx.py [path/to/newsletter_draft_YYYY-MM-DD.md]
  (defaults to the most recent newsletter_draft_*.md in the repo root)
"""

import glob
import os
import re
import sys

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.abspath(__file__))

NAVY = RGBColor(0x1F, 0x38, 0x64)
GREY = RGBColor(0x60, 0x60, 0x60)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = '1F3864'
ZEBRA_FILL = 'EEF1F7'
BODY_FONT = 'Calibri'


def shade_cell(cell, hexfill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), hexfill)
    tcPr.append(shd)


def add_inline(paragraph, text):
    """Add text to a paragraph, honoring **bold** and *italic*."""
    for tok in re.split(r'(\*\*.+?\*\*|\*.+?\*)', text):
        if not tok:
            continue
        if tok.startswith('**') and tok.endswith('**'):
            run = paragraph.add_run(tok[2:-2]); run.bold = True
        elif tok.startswith('*') and tok.endswith('*'):
            run = paragraph.add_run(tok[1:-1]); run.italic = True
        else:
            paragraph.add_run(tok)


def add_table(doc, rows):
    """rows = list of cell-lists; first row is the header."""
    header, body = rows[0], rows[1:]
    table = doc.add_table(rows=1, cols=len(header))
    table.style = 'Table Grid'
    table.alignment = 1
    for j, cell_text in enumerate(header):
        cell = table.rows[0].cells[j]
        shade_cell(cell, HEADER_FILL)
        p = cell.paragraphs[0]
        add_inline(p, cell_text)
        for r in p.runs:
            r.bold = True
            r.font.color.rgb = WHITE
            r.font.size = Pt(9.5)
    for i, row in enumerate(body):
        cells = table.add_row().cells
        for j, cell_text in enumerate(row):
            if j >= len(cells):
                break
            if i % 2 == 1:
                shade_cell(cells[j], ZEBRA_FILL)
            p = cells[j].paragraphs[0]
            add_inline(p, cell_text)
            for r in p.runs:
                r.font.size = Pt(9.5)
    doc.add_paragraph()


def parse_table_row(line):
    return [c.strip() for c in line.strip().strip('|').split('|')]


def is_table_sep(line):
    return bool(re.fullmatch(r'\|?[\s:|-]+\|?', line.strip())) and '-' in line


def build(md_path, out_path):
    lines = open(md_path, encoding='utf-8').read().splitlines()
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = BODY_FONT
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.12
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Inches(0.8)
        section.left_margin = section.right_margin = Inches(0.9)

    i = 0
    while i < len(lines):
        line = lines[i]
        s = line.strip()

        if not s:
            i += 1
            continue

        # Table block
        if s.startswith('|') and i + 1 < len(lines) and is_table_sep(lines[i + 1]):
            rows = [parse_table_row(s)]
            i += 2  # skip header + separator
            while i < len(lines) and lines[i].strip().startswith('|'):
                rows.append(parse_table_row(lines[i]))
                i += 1
            add_table(doc, rows)
            continue

        if s == '---':
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pbdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), 'C0C0C0')
            pbdr.append(bottom); pPr.append(pbdr)
            i += 1
            continue

        if s.startswith('# '):
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(s[2:]); run.bold = True
            run.font.size = Pt(24); run.font.color.rgb = NAVY
            p.paragraph_format.space_after = Pt(2)
        elif s.startswith('### '):
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(s[4:]); run.italic = True
            run.font.size = Pt(10.5); run.font.color.rgb = GREY
            p.paragraph_format.space_after = Pt(10)
        elif s.startswith('## '):
            p = doc.add_paragraph()
            run = p.add_run(s[3:]); run.bold = True
            run.font.size = Pt(15); run.font.color.rgb = NAVY
            p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(4)
        elif s.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            add_inline(p, s[2:])
        else:
            p = doc.add_paragraph()
            add_inline(p, s)
        i += 1

    doc.save(out_path)
    return out_path


def main():
    if len(sys.argv) > 1:
        md_path = sys.argv[1]
    else:
        drafts = sorted(glob.glob(os.path.join(BASE, 'newsletter_draft_*.md')))
        if not drafts:
            sys.exit('No newsletter_draft_*.md found')
        md_path = drafts[-1]
    out_path = os.path.splitext(md_path)[0] + '.docx'
    build(md_path, out_path)
    print(f'Wrote {out_path}')


if __name__ == '__main__':
    main()
